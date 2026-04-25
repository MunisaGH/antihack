"""AI pipeline konkurrent sinovi.

10 ta applicant bir vaqtda ariza yuboradi va ular hammasi tahlildan o'tadimi kuzatadi.
Gemini RPM/quota cheklovlari + thread pool to'g'ri ishlayotganligini tekshiradi.

Ishga tushirish:
    cd backend
    python load_test/ai_flow_test.py

.env da GEMINI_API_KEY va backend ishlayotgan bo'lishi kerak.

Diqqat: Bu real Gemini chaqiruvlari bajaradi — sizning tokenlaringizni sarflaydi!
"""

import concurrent.futures
import io
import os
import random
import string
import sys
import time
import uuid

import requests
from docx import Document

BASE_URL = os.getenv("LOADTEST_URL", "http://localhost:8000")
_RAW_VACANCY_ID = os.getenv("LOADTEST_VACANCY_ID", "")
# Canonical UUID formatiga o'tkazish (defis yoki defissiz qabul qilish)
try:
    VACANCY_ID = str(uuid.UUID(_RAW_VACANCY_ID)) if _RAW_VACANCY_ID else ""
except (ValueError, AttributeError):
    VACANCY_ID = _RAW_VACANCY_ID
NUM_APPLICATIONS = int(os.getenv("LOADTEST_NUM", "5"))


def make_dummy_docx(full_name: str) -> bytes:
    """Haqiqiy rezyume ko'rinishidagi DOCX yaratadi (python-docx orqali)."""
    doc = Document()
    doc.add_heading(full_name, level=0)
    doc.add_paragraph("Email: test@example.com")
    doc.add_paragraph("Phone: +998901234567")
    doc.add_paragraph("Location: Toshkent, O'zbekiston")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Experienced software engineer with 5+ years building scalable web applications. "
        "Strong background in Python, Django, React, and cloud infrastructure. "
        "Passionate about clean code, testing, and performance optimization."
    )

    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Senior Backend Developer — TechCorp (2021-2024)")
    doc.add_paragraph("- Architected microservices handling 10M+ daily requests")
    doc.add_paragraph("- Reduced API latency by 60% through caching and DB optimization")
    doc.add_paragraph("- Led migration from monolith to event-driven architecture")
    doc.add_paragraph("Full-Stack Developer — StartupXYZ (2019-2021)")
    doc.add_paragraph("- Built customer-facing React dashboard with TypeScript")
    doc.add_paragraph("- Integrated OpenAI and Gemini APIs for AI features")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "Python, Django, Django REST Framework, FastAPI, PostgreSQL, Redis, "
        "React, TypeScript, Docker, Kubernetes, AWS, CI/CD, Git, Linux"
    )

    doc.add_heading("Education", level=1)
    doc.add_paragraph("BS Computer Science — Tashkent State University (2015-2019)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def random_phone() -> str:
    return "+998" + "".join(random.choices(string.digits, k=9))


def random_name() -> str:
    first = random.choice(
        ["Ali", "Vali", "Dilshod", "Aziz", "Jamshid", "Otabek", "Nodir", "Saida"]
    )
    last = random.choice(["Olimov", "Karimov", "Usmonov", "Rahimov", "Tursunov"])
    return f"{first} {last}"


def submit_and_track(index: int) -> dict:
    start = time.time()
    result = {
        "index": index,
        "submit_ok": False,
        "analysis_complete": False,
        "final_status": None,
        "score": 0,
        "total_time": 0.0,
        "error": None,
    }

    try:
        name = random_name()
        docx_bytes = make_dummy_docx(name)
        files = {
            "resume_file": (
                "resume.docx",
                io.BytesIO(docx_bytes),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        data = {
            "vacancy": VACANCY_ID,
            "full_name": name,
            "phone": random_phone(),
            "age": str(random.randint(20, 40)),
            "address": "Toshkent",
            "user_language": random.choice(["uz", "ru"]),
        }
        r = requests.post(
            f"{BASE_URL}/api/public/submit_application/",
            data=data,
            files=files,
            timeout=30,
        )
        if r.status_code != 201:
            result["error"] = f"Submit failed: {r.status_code} {r.text[:150]}"
            return result

        app_id = r.json().get("application_id")
        result["submit_ok"] = True
        result["application_id"] = app_id

        # Polling until done
        deadline = start + 120  # 2 daqiqa
        while time.time() < deadline:
            time.sleep(3)
            r2 = requests.get(
                f"{BASE_URL}/api/public/application_status/{app_id}/", timeout=10
            )
            if r2.status_code != 200:
                continue
            data2 = r2.json().get("data", {})
            status = data2.get("status")
            if status and status != "ai_analyzing":
                result["analysis_complete"] = True
                result["final_status"] = status
                result["score"] = data2.get("compatibility_score", 0)
                break
        else:
            result["error"] = "Analysis timeout (>2min)"

    except Exception as e:
        result["error"] = f"Exception: {e}"

    result["total_time"] = round(time.time() - start, 1)
    return result


def main():
    if not VACANCY_ID:
        print("LOADTEST_VACANCY_ID ni o'rnating (mavjud faol vakansiya UUID).")
        print("Misol: $env:LOADTEST_VACANCY_ID='abc-123-def'")
        sys.exit(1)

    print(f"[AI Flow Load Test] {NUM_APPLICATIONS} ta konkurrent ariza")
    print(f"Backend: {BASE_URL}")
    print(f"Vacancy: {VACANCY_ID}\n")

    start = time.time()

    with concurrent.futures.ThreadPoolExecutor(
        max_workers=NUM_APPLICATIONS
    ) as executor:
        futures = [
            executor.submit(submit_and_track, i) for i in range(NUM_APPLICATIONS)
        ]
        results = [f.result() for f in concurrent.futures.as_completed(futures)]

    elapsed = time.time() - start

    # Summary
    submit_ok = sum(1 for r in results if r["submit_ok"])
    analysis_ok = sum(1 for r in results if r["analysis_complete"])
    errors = [r for r in results if r["error"]]

    print("=" * 60)
    print(f"Jami vaqt: {elapsed:.1f}s")
    print(f"Submit muvaffaqiyatli: {submit_ok}/{NUM_APPLICATIONS}")
    print(f"AI tahlili tugagan: {analysis_ok}/{NUM_APPLICATIONS}")
    print(f"Xatoliklar: {len(errors)}")
    print()

    for r in sorted(results, key=lambda x: x["index"]):
        status = "✓" if r["analysis_complete"] else "✗"
        print(
            f"  [{status}] #{r['index']:2d} time={r['total_time']:5.1f}s "
            f"status={r['final_status'] or 'N/A':20s} score={r['score']:3d}  "
            f"{r['error'] or ''}"
        )


if __name__ == "__main__":
    main()
